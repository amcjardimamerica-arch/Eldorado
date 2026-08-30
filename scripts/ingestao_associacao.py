"""Indexa DOCX de uma associação, removendo PII da cópia pública compactada."""
from __future__ import annotations
import gzip, hashlib, json, re, sys
from pathlib import Path
from zipfile import ZipFile
sys.path.insert(0,str(Path(__file__).resolve().parents[1]))
from src.nucleo import ROOT, has_prompt_injection, now_iso, slug
from scripts.criar_associacao import create

CPF=re.compile(r"\b\d{3}\.?\d{3}\.?\d{3}-?\d{2}\b")
RG_LINE=re.compile(r"(?i)\bRG\b\s*[:|]?\s*[0-9.\-]+")
PERSONAL_EMAIL=re.compile(r"(?i)\b(?!amc\.jardimamerica@gmail\.com)[\w.+-]+@(gmail|hotmail|outlook|yahoo)\.[\w.-]+\b")

def sanitize(text: str) -> str:
    return PERSONAL_EMAIL.sub("[E-MAIL PESSOAL REDIGIDO]",RG_LINE.sub("RG [REDIGIDO]",CPF.sub("[CPF REDIGIDO]",text)))

def sanitize_tables(tables):
    safe=json.loads(sanitize(json.dumps(tables,ensure_ascii=False)))
    for table in safe:
        for row in table:
            for index,value in enumerate(row[:-1]):
                if str(value).strip().lower() in {"rg","cpf"}: row[index+1]="[REDIGIDO]"
    return safe

def docx_content(path: Path):
    from docx import Document
    from PIL import Image
    doc=Document(path); paragraphs=[p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables=[]
    for table in doc.tables:
        rows=[]
        for row in table.rows:
            values=[]
            for cell in row.cells:
                value=" ".join(cell.text.split())
                if not values or value!=values[-1]: values.append(value)
            rows.append(values)
        tables.append(rows)
    media=[]
    with ZipFile(path) as z:
        for name in z.namelist():
            if not name.startswith("word/media/"): continue
            raw=z.read(name); dimensions=None
            try:
                import io
                with Image.open(io.BytesIO(raw)) as im: dimensions=[im.width,im.height]
            except Exception: pass
            media.append({"nome_interno":Path(name).name,"sha256":hashlib.sha256(raw).hexdigest(),"bytes":len(raw),"dimensoes_px":dimensions})
    return paragraphs,tables,media

def ingest(aid: str, file: Path) -> dict:
    root=create(aid,aid); raw=file.read_bytes(); digest=hashlib.sha256(raw).hexdigest()
    if file.suffix.lower()!=".docx": raise ValueError("esta versão aceita DOCX; outros formatos permanecem no armazenamento privado")
    paragraphs,tables,media=docx_content(file)
    joined="\n\n".join(paragraphs); injection=has_prompt_injection(joined)
    if injection: raise ValueError("documento em quarentena por possível prompt injection")
    safe_text=sanitize(joined); safe_tables=sanitize_tables(tables)
    public=root/"documentos"; public.mkdir(parents=True,exist_ok=True)
    (public/(digest+".texto.md.gz")).write_bytes(gzip.compress(safe_text.encode(),compresslevel=9,mtime=0))
    (public/(digest+".tabelas.json.gz")).write_bytes(gzip.compress(json.dumps(safe_tables,ensure_ascii=False,separators=(",",":")).encode(),compresslevel=9,mtime=0))
    manifest={"arquivo_original":file.name,"sha256":digest,"bytes":len(raw),"indexado_em":now_iso(),"prompt_injection":False,"paragrafos":len(paragraphs),"tabelas":len(tables),"imagens":len(media),"pii_redigida":["cpf","rg","email_pessoal"],"original_publicado":False,"artefatos_publicos":[digest+".texto.md.gz",digest+".tabelas.json.gz"]}
    (public/"manifesto_publico.json").write_text(json.dumps(manifest,ensure_ascii=False,indent=2)+"\n",encoding="utf-8")
    (root/"imagens"/"manifesto_publico.json").write_text(json.dumps({"documento_sha256":digest,"originais_publicados":False,"itens":media},ensure_ascii=False,indent=2)+"\n",encoding="utf-8")
    return manifest

if __name__=="__main__":
    if len(sys.argv)!=3: raise SystemExit("uso: ingestao_associacao.py ASSOCIACAO arquivo.docx")
    print(json.dumps(ingest(slug(sys.argv[1]),Path(sys.argv[2])),ensure_ascii=False,indent=2))
